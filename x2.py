#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue May 14 19:42:09 2019

@author: tkhamsi

Goal 1: get running at all - Success!
Goal 2: store, reapply formatting
3: if statement regarding if there's a file w that name already; oh, rn it just overwrites the existing one
4: small text doc storing words that this program changed
5: exception for xxx already, with an admonishing message (we mustn't be greedy)
        OR change the title to 'you thought you could play me?'
        orrrr, replace the text with 'nice try.' all the text. 
Goal 2: do so for 'cks', 'chs' as well
Goal 3: with stupid hand-drawn Xs
Goal 4: PDFs (professional ones). i'll do it to mine.
"""

import docx

def getText(filename):
    doc = docx.Document(filename)

#index = 0
#lines = len(doc.paragraphs)

    listmin = []
    for paragraph in doc.paragraphs:
        listmin.append(paragraph.text)
    return (listmin)

#finds x indexes, replaces w triple x. this may be too much. we shall see. 
#do i split the line containing x? this is short, so let's try shit. 
def coolFile(textList):
    coolList = []
    x = 'xxx'
    for line in textList:
        if 'x' in line:
            l = line.split('x') #makes a list l containing each half
            l2 = l[0] + x + l[1]
            coolList.append(l2)
        else:
            coolList.append(line)
    return coolList

def writeFile(coolList, filename): #it's fine that this name here, ya?
    newDoc = docx.Document()
    for line in coolList:
        newDoc.add_paragraph(line)
    newDoc.save('Cool ' + filename)
    return newDoc

filename = 'x2test3.docx' #eventually I'll make this a user query
testme = getText(filename)
test2 = coolFile(testme)
test3 = writeFile(test2, filename)
print (test3)
#print (test2) 




#print (testme)


"""
with open('x2test2.docx', 'r', encoding = 'latin-1') as f:
    for line in f:
        print (line)
"""
